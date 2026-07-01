import os
import re
import shutil
import sys
import time

import win32com.client
import win32com.client.gencache
import pandas as pd

from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import core.constants as C
from core.scanning import extract_dia_code


def getRevisionInteger(section, strRev, newRev):
    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Revision" in cell.text:
                    strRev = cell.text[C._LEN_REVISION_STR:]

    if strRev.isalpha():
        newRev = chr(ord(strRev) + 1)
    else:
        newRev = int(strRev) + 1
        if newRev < 10:
            newRev = '0' + str(newRev)
        else:
            newRev = str(newRev)

    return newRev


def headerRevUpdate(docObject, strRev, newRev):
    section = docObject.sections[0]
    newRev  = getRevisionInteger(section, strRev, newRev)

    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Revision" in cell.text:
                    cell.text = "Revision No: " + newRev
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return newRev


def revHistoryUpdate(docObject, intRev, revHistoryTableCommentary, strVersion):
    listTables = docObject.tables
    revHistoryTable = listTables[-1] if listTables else None
    newRow = revHistoryTable.add_row()
    newRow.cells[0].text = str(intRev)
    newRow.cells[1].text = revHistoryTableCommentary + strVersion


def docPathRevUpdate(docObject, oldDocPath, intRev, newDocPath):
    whitespaceMatch = re.search(C.REGEX_DOC_PATH_REV_WHITESPACE, oldDocPath)
    if whitespaceMatch:
        newDocPath = re.sub(C.REGEX_DOC_PATH_REV_WHITESPACE, ("Rev" + str(intRev)), oldDocPath)
    else:
        newDocPath = re.sub(C.REGEX_DOC_PATH_REV, ("Rev" + str(intRev)), oldDocPath)
    docObject.save(newDocPath)
    return newDocPath


def handleDispatch(app_name):
    try:
        app = win32com.client.gencache.EnsureDispatch(app_name)
    except AttributeError:
        MODULE_LIST = [m.__name__ for m in sys.modules.values()]
        for module in MODULE_LIST:
            if re.match(r'win32com\.gen_py\..+', module):
                del sys.modules[module]
        shutil.rmtree(win32com.client.gencache.GetGeneratePath())
        app = win32com.client.gencache.EnsureDispatch(app_name)
    return app


def _wait_for_word_exit(timeout=30):
    import subprocess
    for _ in range(timeout):
        result = subprocess.run(
            ["tasklist", "/FI", "IMAGENAME eq WINWORD.EXE", "/NH"],
            capture_output=True, text=True,
            creationflags=subprocess.CREATE_NO_WINDOW
        )
        if "WINWORD.EXE" not in result.stdout:
            return
        time.sleep(1)


def createWordComparison(oldDocPath, newDocPath):
    from urllib.parse import unquote
    oldDocPath = os.path.normpath(unquote(str(oldDocPath)))
    newDocPath = os.path.normpath(unquote(str(newDocPath)))
    numRetries = 0
    while numRetries < C._MAX_RETRIES:
        _wait_for_word_exit()
        wordApp = None
        try:
            wordApp = handleDispatch("Word.Application")
            wordApp.Visible = False
            old_doc = wordApp.Documents.Open(oldDocPath, Visible=False)
            new_doc = wordApp.Documents.Open(newDocPath, Visible=False)
            wordApp.CompareDocuments(old_doc, new_doc)
        except Exception as e:
            print(f"[DEBUG] Attempt {numRetries + 1} failed: {e}")
            numRetries += 1
            if wordApp:
                try:
                    wordApp.Quit()
                except Exception:
                    pass
                del wordApp
                _wait_for_word_exit()
            time.sleep(C._DELAY * numRetries)
        else:
            try:
                wordApp.ActiveDocument.SaveAs(FileName=(oldDocPath[:(len(oldDocPath) - 5)] + "_REDLINE.docx"))
                old_doc.Close(False)
                new_doc.Close(False)
                wordApp.Quit()
            except Exception:
                try:
                    wordApp.Quit()
                except Exception:
                    pass
            del wordApp
            _wait_for_word_exit()
            return

    raise Exception("Max retries exceeded. Comparison function not run.")


def extract_revision_from_filename(filename):
    match = re.search(r'Rev\s?([A-Z0-9]+)', filename, re.IGNORECASE)
    return match.group(1) if match else ""


def revision_sort_key(rev_str):
    if rev_str.isalpha():
        return (0, rev_str)
    try:
        return (1, int(rev_str))
    except ValueError:
        return (2, rev_str)


def find_top_two_revisions(docx_path):
    docx_basename = os.path.basename(docx_path)
    item_match    = re.search(r"10\d{4}", docx_basename)
    if not item_match:
        return None, None

    item_no       = item_match.group(0)
    matching_files = []
    for root, _, files in os.walk(C.RECORDS_FOLDER):
        for fname in files:
            if fname.lower().endswith('.docx') and item_no in fname and 'red' not in fname.lower():
                full_path = os.path.join(root, fname)
                rev = extract_revision_from_filename(fname)
                matching_files.append((rev, full_path, fname))

    if len(matching_files) < 2:
        return None, None

    matching_files.sort(key=lambda x: revision_sort_key(x[0]))
    return matching_files[-2][1], matching_files[-1][1]


def remove_empty_pages(doc):
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)
        doc.paragraphs.pop()


def find_date(firstPageStr):
    date_pattern = re.compile(
        r'\b(\d{1,2}/\d{1,2}/\d{4}|(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b',
        re.IGNORECASE
    )
    dates = date_pattern.findall(firstPageStr)
    return re.sub(r'\s+', ' ', dates[0][0]).strip()


def read_word_document(pdf_path, docx_path):
    pdfReader    = PdfReader(pdf_path)
    doc          = Document(docx_path)
    firstPageStr = pdfReader.pages[0].extract_text()

    versionPos = firstPageStr.find("Version:") if firstPageStr.find("Version:") > -1 else firstPageStr.find("Version")
    versionStr = firstPageStr[(versionPos + C._LEN_VERSION_STR):(versionPos + C._LEN_END_VERSION_STR)].strip()
    dateStr    = find_date(firstPageStr)

    return doc, versionStr, dateStr


def process_documents(docDF, mode, doc_indices=None, progress_callback=None):
    if mode in (1, 3):
        os.makedirs(C.REDLINE_FOLDER, exist_ok=True)
    if mode in (1, 2):
        os.makedirs(C.PREV_REV_FOLDER, exist_ok=True)

    rows_to_process = doc_indices if doc_indices is not None else range(len(docDF))
    total      = len(rows_to_process)
    pdf_paths  = docDF["PDF Path"].values
    docx_paths = docDF["DOCX Path"].values

    print(f"Path: {C.RECORDS_FOLDER}\n")

    for i, row_idx in enumerate(rows_to_process):
        if progress_callback:
            progress_callback(i + 1, total)

        pdf            = pdf_paths[row_idx]
        docPathInitial = docx_paths[row_idx]

        if mode in (1, 2):
            if not pdf or (isinstance(pdf, float) and pd.isna(pdf)) or not str(pdf).strip():
                print(f"[DEBUG] Skipped: Row {row_idx + 1} — Missing PDF path")
                continue
        if not docPathInitial or (isinstance(docPathInitial, float) and pd.isna(docPathInitial)) or not str(docPathInitial).strip():
            print(f"[DEBUG] Skipped: Row {row_idx + 1} — Missing DOCX path")
            continue

        print(f"\n{'='*60}")
        print(f"Processing: {os.path.basename(docPathInitial)}")

        if mode == 3:
            oldDocPath, newDocPath = find_top_two_revisions(docPathInitial)
            if oldDocPath is None or newDocPath is None:
                print("  [SKIPPED] Document has fewer than 2 versions. Cannot create redline.")
                continue

            print(f"  Comparing: {os.path.basename(oldDocPath)}")
            print(f"        vs:  {os.path.basename(newDocPath)}")
            createWordComparison(oldDocPath, newDocPath)

            redline_path = oldDocPath[:(len(oldDocPath) - 5)] + "_REDLINE.docx"
            if os.path.exists(redline_path):
                redline_filename  = os.path.basename(redline_path)
                dest_redline_path = os.path.join(C.REDLINE_FOLDER, redline_filename)
                shutil.move(redline_path, dest_redline_path)
                print(f"Moved redline to Redline folder:{redline_filename}")
            continue

        revisionStr = ""
        revisionInt = 0
        docPathNew  = ""

        doc, versionStr, dateStr = read_word_document(pdf, docPathInitial)

        revisionInt  = headerRevUpdate(doc, revisionStr, revisionInt)
        textToUpdate = ""
        for para in doc.paragraphs:
            if "DIA" in para.text:
                textToUpdate = para.text

        textToUpdate = re.sub(C.REGEX_REVISION_PATTERN, versionStr, str(textToUpdate))
        textToUpdate = re.sub(C.REGEX_DATE_PATTERN,     dateStr,    str(textToUpdate))

        for para in doc.paragraphs:
            if "DIA" in para.text:
                para.text = textToUpdate

        revHistoryUpdate(doc, revisionInt, C.STR_REV_COMMENTARY, C.STR_VERSION)
        remove_empty_pages(doc)
        docPathNew = docPathRevUpdate(doc, docPathInitial, revisionInt, docPathNew)

        pdf_label = os.path.splitext(os.path.basename(pdf))[0]
        print(f"\nCreated: {os.path.basename(docPathNew)}")
        print(f"         [{pdf_label} – Date: {dateStr}]")

        if mode == 1:
            createWordComparison(docPathInitial, docPathNew)
            redline_path = docPathInitial[:(len(docPathInitial) - 5)] + "_REDLINE.docx"
            if os.path.exists(redline_path):
                dest_redline = os.path.join(C.REDLINE_FOLDER, os.path.basename(redline_path))
                shutil.move(redline_path, dest_redline)
                print(f"Moved redline to Redline folder:{os.path.basename(redline_path)}")
        else:
            print("[DEBUG] Skipped Redline: As requested for 'Create only new rev' mode")

        docx_filename = os.path.splitext(os.path.basename(docPathInitial))[0]
        item_match    = re.search(r"10\d{4}", docx_filename)
        text_after    = docx_filename[item_match.end():] if item_match else docx_filename
        text_after    = re.sub(r"^[,\s]+", "", text_after)
        text_after    = re.sub(r",?\s*Rev\s?[A-Z0-9]+", "", text_after).strip().rstrip(",").strip()

        dia_id        = extract_dia_code(os.path.basename(pdf)) or ""
        part_number   = item_match.group(0) if item_match else ""
        folder_prefix = f"{part_number}, " if part_number else ""
        dest_folder   = os.path.join(C.RECORDS_FOLDER, f"{folder_prefix}{text_after} ({dia_id})")

        if os.path.exists(docPathInitial) and not re.search(r'red', os.path.basename(docPathInitial), re.IGNORECASE):
            shutil.move(docPathInitial, os.path.join(C.PREV_REV_FOLDER, os.path.basename(docPathInitial)))
        if os.path.exists(docPathNew) and not re.search(r'red', os.path.basename(docPathNew), re.IGNORECASE):
            os.makedirs(dest_folder, exist_ok=True)
            shutil.move(docPathNew, os.path.join(dest_folder, os.path.basename(docPathNew)))
        if os.path.exists(pdf):
            os.makedirs(dest_folder, exist_ok=True)
            shutil.move(pdf, os.path.join(dest_folder, os.path.basename(pdf)))

        print(f"[DEBUG] Organized: {folder_prefix}{text_after} ({dia_id})")
