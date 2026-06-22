import json
import os
import re
import shutil
import sys
import time

import openpyxl
import pandas as pd

from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import win32com.client
import win32com.client.gencache

alphabet = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J',
            'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T',
            'U', 'V', 'W', 'X', 'Y', 'Z']

__lenRevisionStr = 13
__maxRetries = 3
__delay = 1

regexDocPathRevOnItemNoPattern = r"10\d\d\d\d-\d\d"
regexDocPathRevWhitespacePattern = r"Rev\s[A-Z0-9]+"
regexDocPathItemNoPattern = r"10\d\d\d\d"
regexDocPathRev = r"Rev[A-Z0-9]+"


def getRevisionInteger(section, strRev, newRev):
    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Revision" in cell.text:
                    strRev = cell.text[__lenRevisionStr:]

    if strRev.isalpha():
        posRev = alphabet.index(strRev)
        newPosRev = posRev + 1
        newRev = alphabet[newPosRev]
    else:
        newRev = int(strRev) + 1
        if newRev < 10:
            newRev = '0' + str(newRev)
        else:
            newRev = str(newRev)

    return newRev


def headerRevUpdate(docObject, strRev, newRev):
    section = docObject.sections[0]
    newRev = getRevisionInteger(section, strRev, newRev)

    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Revision" in cell.text:
                    cell.text = "Revision No: " + newRev
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return newRev


def revHistoryUpdate(docObject, intRev, revHistoryTableCommentary, strVersion):
    listTables = docObject.tables
    if listTables:
        revHistoryTable = listTables[-1]
    else:
        revHistoryTable = None

    newRow = revHistoryTable.add_row()
    newRow.cells[0].text = str(intRev)
    newRow.cells[1].text = revHistoryTableCommentary + strVersion


def docPathRevUpdate(docObject, oldDocPath, intRev, newDocPath):
    revEndFileNameWhitespaceMatch = re.search(regexDocPathRevWhitespacePattern, oldDocPath)

    if revEndFileNameWhitespaceMatch:
        newDocPath = re.sub(regexDocPathRevWhitespacePattern, ("Rev" + str(intRev)), oldDocPath)
    else:
        newDocPath = re.sub(regexDocPathRev, ("Rev" + str(intRev)), oldDocPath)

    docObject.save(newDocPath)
    return newDocPath


def handleDispatch(app_name):
    try:
        app = win32com.client.gencache.EnsureDispatch(app_name)
    except AttributeError:
        import sys
        import shutil
        MODULE_LIST = [m.__name__ for m in sys.modules.values()]
        for module in MODULE_LIST:
            if re.match(r'win32com\.gen_py\..+', module):
                del sys.modules[module]
        shutil.rmtree(win32com.client.gencache.GetGeneratePath())
        app = win32com.client.gencache.EnsureDispatch(app_name)
    return app


def createWordComparison(oldDocPath, newDocPath):
    numRetries = 0

    while numRetries < __maxRetries:
        try:
            wordApp = handleDispatch("Word.Application")
            wordApp.Visible = False
            old_doc = wordApp.Documents.Open(oldDocPath)
            new_doc = wordApp.Documents.Open(newDocPath)
            wordApp.CompareDocuments(old_doc, new_doc)

        except Exception as e:
            print(f"Attempt {numRetries + 1} failed: {e}")
            numRetries += 1
            time.sleep(__delay * numRetries)

        else:
            wordApp.ActiveDocument.SaveAs(FileName = (oldDocPath[:(len(oldDocPath) - 5)] + "_REDLINE.docx"))
            wordApp.Quit()
            return

    raise Exception("Max retries exceeded. Comparison function not run.")

#################### MACROS ####################
_SCRIPT_DIR     = os.path.dirname(os.path.abspath(__file__))
RECORDS_FOLDER  = os.path.join(_SCRIPT_DIR, 'KBM Docs')
OUTDATED_FOLDER = os.path.join(RECORDS_FOLDER, 'outdated')
JSON_PATH       = os.path.join(_SCRIPT_DIR, 'kbm_map.json')

__lenVersionStr = 9
__lenEndVersionStr = 14

regexRevisionPattern = '[0-9]+\.[0-9]+'
regexDatePattern = '[0-9]+/[0-9]+/[0-9]+'
regexDocPathItemNoPattern = r"[0-9]+"
regexDocPathRevOnItemNoPattern = r"[0-9]+-[0-9]+"
regexDocPathRevWhitespacePattern = r"[A-Za-z]+\s[0-9]+"
regexDocPathRev = r"[A-Za-z]+[0-9]+"

regexDocPathRevOnItemNoPattern_dia = r"10\d\d\d\d-\d\d"
regexDocPathRevWhitespacePattern_dia = r"Rev\s[A-Z0-9]+"
regexDocPathItemNoPattern_dia = r"10\d\d\d\d"
regexDocPathRev_dia = r"Rev[A-Z0-9]+"

strRevCommentary = "External document updated."
strVersion = ""

#################### END Macros ####################

#################### LOCAL VARIABLES ####################
revisionStr = ""
revisionInt = 0
docPathNew = ""
#################### END VARIABLES ####################

# ═══════════════════════════════════════════════════════════════════════════════
# EXCEL UPDATE FUNCTIONS (from update_whitney_excel.py)
# ═══════════════════════════════════════════════════════════════════════════════

def load_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def normalize(s):
    return s.replace('_', ' ')


def extract_dia_code(filename):
    normalized = normalize(filename)
    match = re.search(r'\bDIA\s+([A-Z0-9][A-Z0-9\-\.]*[A-Z0-9]|[A-Z]+)', normalized)
    if match:
        return f"DIA {match.group(1)}"
    return None


def get_id_from_json(data, dia_code, normalized_filename=None):
    entry = data.get(dia_code)
    if entry is not None:
        return entry if isinstance(entry, str) else entry.get('id')
    if normalized_filename:
        for key, value in data.items():
            if key in normalized_filename:
                return value if isinstance(value, str) else value.get('id')
    return None


def version_score(filename):
    m = re.search(r'\bv\.?(\d+)(?:\.(\d+))?', filename, re.IGNORECASE)
    if m:
        major = int(m.group(1))
        minor = int(m.group(2)) if m.group(2) else 0
        return (major, minor)
    return (0, 0)


def collect_files(folder):
    pdf_files  = []
    docx_files = []
    skip = {'outdated', 'redline'}
    for root, dirs, files in os.walk(folder):
        dirs[:] = [d for d in dirs if d.lower() not in skip]
        for name in files:
            lower = name.lower()
            full  = os.path.abspath(os.path.join(root, name))
            if lower.endswith('.pdf'):
                pdf_files.append(full)
            elif lower.endswith('.docx'):
                docx_files.append(full)
    return pdf_files, docx_files


def build_docx_index(docx_files):
    index = {}
    pattern = re.compile(r'^(\d{6}(?:-\d{3,4})?)\b')
    for path in docx_files:
        name = os.path.basename(path)
        m = pattern.match(name)
        if m:
            prefix = m.group(1)
            index.setdefault(prefix, []).append(path)
    return index


def create_workbook():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['Index', 'PDF Path', 'DOCX Path'])
    ws.column_dimensions['B'].width = 124
    ws.column_dimensions['C'].width = 124
    return wb, ws


def move_outdated_pdfs(all_pdfs_by_id, best_pdf):
    moved = 0
    os.makedirs(OUTDATED_FOLDER, exist_ok=True)
    for doc_id, entries in all_pdfs_by_id.items():
        if len(entries) <= 1:
            continue
        best_path = best_pdf[doc_id][1]
        for _, pdf_path in entries:
            if pdf_path == best_path:
                continue
            dest = os.path.join(OUTDATED_FOLDER, os.path.basename(pdf_path))
            shutil.move(pdf_path, dest)
            print(f'  [outdated] moved: {os.path.basename(pdf_path)}')
            moved += 1
    return moved


def update_excel(excel_path):
    """Run the excel update logic and return list of missing doc IDs"""
    data       = load_json(JSON_PATH)
    pdf_files, docx_files = collect_files(RECORDS_FOLDER)

    docx_files = [p for p in docx_files if not re.search(r'red', os.path.basename(p), re.IGNORECASE)]
    docx_index = build_docx_index(docx_files)
    wb, ws     = create_workbook()
    start_row  = 2
    index      = 1

    no_dia  = 0
    no_json = 0

    all_pdfs_by_id = {}
    for pdf_path in pdf_files:
        name     = os.path.basename(pdf_path)
        dia_code = extract_dia_code(name)
        if dia_code is None:
            no_dia += 1
            continue
        doc_id = get_id_from_json(data, dia_code, normalize(name))
        if doc_id is None:
            no_json += 1
            print(f'  [no JSON match]  {name}  →  {dia_code}')
            continue
        score = version_score(name)
        all_pdfs_by_id.setdefault(doc_id, []).append((score, pdf_path))

    best_pdf = {
        doc_id: max(entries, key=lambda x: x[0])
        for doc_id, entries in all_pdfs_by_id.items()
    }

    moved_count = move_outdated_pdfs(all_pdfs_by_id, best_pdf)

    no_docx        = 0
    matched        = 0
    missing_doc_ids = []

    for doc_id, (score, pdf_path) in sorted(best_pdf.items()):
        candidates = docx_index.get(doc_id, [])
        docx_path  = candidates[0] if candidates else None

        if docx_path is None:
            no_docx += 1
            missing_doc_ids.append(doc_id)

        ws.cell(row=start_row, column=1, value=index)
        ws.cell(row=start_row, column=2, value=pdf_path)
        ws.cell(row=start_row, column=3, value=docx_path if docx_path else '')

        index     += 1
        start_row += 1
        matched   += 1

    wb.save(excel_path)

    print(f'\nDone. Wrote {matched} rows to {excel_path}')
    print(f'  Outdated PDFs moved to outdated/: {moved_count}')
    print(f'  PDFs without DIA identifier:       {no_dia}')
    print(f'  DIA codes not in JSON:             {no_json}')
    print(f'  DIA codes with no DOCX match:      {no_docx}')

    if missing_doc_ids:
        print('\nDOCX files needed (add to folder):')
        for doc_id in missing_doc_ids:
            print(f'  {doc_id}')

    return missing_doc_ids


# ═══════════════════════════════════════════════════════════════════════════════
# MISSING DOCS MENU
# ═══════════════════════════════════════════════════════════════════════════════

def get_user_menu_choice(prompt, options):
    print(f"\n{prompt}")
    for i, option in enumerate(options, 1):
        print(f"[{i}] {option}")
    while True:
        try:
            choice = int(input("\nEnter your choice: ").strip())
            if 1 <= choice <= len(options):
                return choice
            print(f"Please enter a number between 1 and {len(options)}")
        except ValueError:
            print("Invalid input. Please enter a number.")


def handle_missing_docs():
    """Show menu for missing docs and return True if user wants to continue"""
    choice = get_user_menu_choice(
        "Missing DOCX files detected!",
        [
            "Abort and fix KBM Docs",
            "Keep going (add to excel but skip processing)"
        ]
    )
    return choice == 2


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT PROCESSING FUNCTIONS (from whitney_doc_tool.py)
# ═══════════════════════════════════════════════════════════════════════════════

def extract_revision_from_filename(filename):
    match = re.search(r'Rev\s?([A-Z0-9]+)', filename, re.IGNORECASE)
    return match.group(1) if match else ""


def revision_sort_key(rev_str):
    if rev_str.isalpha():
        return (0, rev_str)
    else:
        try:
            return (1, int(rev_str))
        except ValueError:
            return (2, rev_str)


def find_top_two_revisions(docx_path):
    docx_dir = os.path.dirname(docx_path)
    docx_basename = os.path.basename(docx_path)
    parent_dir = os.path.dirname(docx_dir)

    item_match = re.search(r"10\d{4}", docx_basename)
    if not item_match:
        return None, None

    item_no = item_match.group(0)

    matching_files = []
    for root, _, files in os.walk(parent_dir):
        for fname in files:
            if fname.lower().endswith('.docx') and item_no in fname and 'red' not in fname.lower():
                full_path = os.path.join(root, fname)
                rev = extract_revision_from_filename(fname)
                matching_files.append((rev, full_path, fname))

    if len(matching_files) < 2:
        return None, None

    matching_files.sort(key=lambda x: revision_sort_key(x[0]))

    return matching_files[-2][1], matching_files[-1][1]


def select_documents(pdf_paths):
    print("\nAvailable PDFs:")
    for i, pdf_path in enumerate(pdf_paths, 1):
        basename = os.path.basename(pdf_path)
        print(f"[{i}] {basename}")

    while True:
        try:
            user_input = input("\nEnter PDF numbers to process (comma-separated, e.g., 1,3,5): ").strip()
            indices = [int(x.strip()) - 1 for x in user_input.split(',')]
            if all(0 <= idx < len(pdf_paths) for idx in indices):
                return indices
            print(f"Please enter valid numbers between 1 and {len(pdf_paths)}")
        except ValueError:
            print("Invalid input. Please enter numbers separated by commas (e.g., 1,2,3)")


def remove_empty_pages(doc):
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)
        doc.paragraphs.pop()


def find_date(firstPageStr):
    date_pattern = re.compile(
        r'\b(\d{1,2}/\d{1,2}/\d{4}|(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b',
        re.IGNORECASE
    )
    dates = date_pattern.findall(firstPageStr)
    return re.sub(r'\s+', ' ', dates[0][0]).strip()


def read_word_document(pdf_path, docx_path):
    pdfReader = PdfReader(pdf_path)
    doc = Document(docx_path)
    page = pdfReader.pages[0]
    firstPageStr = page.extract_text()

    versionPos = firstPageStr.find("Version:") if firstPageStr.find("Version:") > -1 else firstPageStr.find("Version")
    versionStr = firstPageStr[(versionPos + __lenVersionStr):(versionPos + __lenEndVersionStr)].strip()
    dateStr     = find_date(firstPageStr)

    print(f"{docx_path}|{versionStr}|{dateStr}")
    return doc, versionStr, dateStr


def process_documents(docDF, mode, doc_indices=None):
    """Process documents based on selected mode and scope"""
    global revisionStr, revisionInt, docPathNew

    if doc_indices is None:
        rows_to_process = range(len(docDF))
    else:
        rows_to_process = doc_indices

    pdf_paths = docDF["PDF Path"].values
    docx_paths = docDF["DOCX Path"].values

    for row_idx in rows_to_process:
        pdf = pdf_paths[row_idx]
        docPathInitial = docx_paths[row_idx]

        # Skip rows with missing PDF or DOCX paths
        if not pdf or (isinstance(pdf, float) and pd.isna(pdf)) or not str(pdf).strip():
            print(f"\n[SKIPPED] Row {row_idx + 1}: Missing PDF path")
            continue
        if not docPathInitial or (isinstance(docPathInitial, float) and pd.isna(docPathInitial)) or not str(docPathInitial).strip():
            print(f"\n[SKIPPED] Row {row_idx + 1}: Missing DOCX path")
            continue

        print(f"\n{'='*60}")
        print(f"Processing row {row_idx + 1}: {os.path.basename(docPathInitial)}")
        print(f"{'='*60}")

        # Mode 3: Create only redline (compare two existing versions)
        if mode == 3:
            oldDocPath, newDocPath = find_top_two_revisions(docPathInitial)
            if oldDocPath is None or newDocPath is None:
                print(f"  [SKIPPED] Document has fewer than 2 versions. Cannot create redline.")
                continue

            print(f"  Comparing: {os.path.basename(oldDocPath)}")
            print(f"        vs:  {os.path.basename(newDocPath)}")
            createWordComparison(oldDocPath, newDocPath)
            time.sleep(2)

            redline_path = oldDocPath[:(len(oldDocPath) - 5)] + "_REDLINE.docx"
            if os.path.exists(redline_path):
                redline_filename = os.path.basename(redline_path)
                dest_redline_path = os.path.join(RECORDS_FOLDER, redline_filename)
                if os.path.dirname(redline_path) != RECORDS_FOLDER:
                    shutil.move(redline_path, dest_redline_path)
                    print(f"  Moved redline to KBM Docs folder: {redline_filename}")
                else:
                    print(f"  Created redline: {redline_filename}")

            continue

        # Mode 1 & 2: Create new revision (and possibly redline)
        doc, versionStr, dateStr = read_word_document(pdf, docPathInitial)

        revisionInt = headerRevUpdate(doc, revisionStr, revisionInt)

        for para in doc.paragraphs:
            if "DIA" in para.text:
                textToUpdate = para.text

        textToUpdate = re.sub(regexRevisionPattern, versionStr, str(textToUpdate))
        textToUpdate = re.sub(regexDatePattern, dateStr, str(textToUpdate))

        for para in doc.paragraphs:
            if "DIA" in para.text:
                para.text = textToUpdate

        revHistoryUpdate(doc, revisionInt, strRevCommentary, strVersion)
        remove_empty_pages(doc)

        docPathNew = docPathRevUpdate(doc, docPathInitial, revisionInt, docPathNew)

        if mode == 1:
            createWordComparison(docPathInitial, docPathNew)
            time.sleep(2)
        else:  # mode == 2
            print(f"  [SKIPPED REDLINE] As requested for 'Create only new rev' mode")

        docx_dir = os.path.dirname(docPathInitial)
        docx_filename = os.path.splitext(os.path.basename(docPathInitial))[0]

        item_match = re.search(r"10\d{4}", docx_filename)
        text_after = docx_filename[item_match.end():] if item_match else docx_filename
        text_after = re.sub(r"^[,\s]+", "", text_after)
        text_after = re.sub(r",?\s*Rev\s?[A-Z0-9]+", "", text_after).strip().rstrip(",").strip()

        dia_match = re.search(r"DIA\s+[A-Z]+-\d+[A-Z]?\.\d+", os.path.basename(pdf))
        dia_id = dia_match.group(0) if dia_match else ""

        dest_folder = os.path.join(docx_dir, f"{text_after} ({dia_id})")
        os.makedirs(dest_folder, exist_ok=True)

        if os.path.exists(docPathInitial) and not re.search(r'red', os.path.basename(docPathInitial), re.IGNORECASE):
            shutil.move(docPathInitial, os.path.join(dest_folder, os.path.basename(docPathInitial)))
        if os.path.exists(docPathNew) and not re.search(r'red', os.path.basename(docPathNew), re.IGNORECASE):
            shutil.move(docPathNew, os.path.join(dest_folder, os.path.basename(docPathNew)))
        if os.path.exists(pdf):
            shutil.move(pdf, os.path.join(dest_folder, os.path.basename(pdf)))

        print(f"  Organized: {text_after} ({dia_id})")


def main():
    if len(sys.argv) < 2:
        print("Usage: python whitney_together.py <excel_filename>")
        print("Example: python whitney_together.py \"whitney testing.xlsx\"")
        sys.exit(1)

    excel_arg  = sys.argv[1]
    EXCEL_PATH = excel_arg if os.path.isabs(excel_arg) else os.path.join(_SCRIPT_DIR, excel_arg)

    # Step 1: Update the excel file
    print("="*60)
    print("STEP 1: Updating Excel with current KBM Docs...")
    print("="*60)
    missing_doc_ids = update_excel(EXCEL_PATH)

    # Step 2: Check for missing docs
    if missing_doc_ids:
        if not handle_missing_docs():
            print("\nAborted. Please update KBM Docs and run again.")
            sys.exit(0)
        print("\nContinuing with processing (missing docs will be skipped)...")

    # Step 3: Read the excel file
    print("\n" + "="*60)
    print("STEP 2: Processing Documents...")
    print("="*60)
    docDF = pd.read_excel(EXCEL_PATH)

    # First menu: Choose operation mode
    mode = get_user_menu_choice(
        "Select the operation to perform:",
        [
            "Create new rev and redline",
            "Create only new rev",
            "Create only new redline"
        ]
    )

    # Second menu: Choose scope
    scope = get_user_menu_choice(
        "Select scope:",
        [
            "Affect all documents",
            "Select docs"
        ]
    )

    doc_indices = None
    if scope == 2:
        pdf_list = docDF["PDF Path"].values.tolist()
        doc_indices = select_documents(pdf_list)

    # Process documents
    process_documents(docDF, mode, doc_indices)

    print(f"\n{'='*60}")
    print("All operations completed!")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
