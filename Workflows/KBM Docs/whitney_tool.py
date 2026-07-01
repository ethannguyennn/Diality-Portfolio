import warnings
warnings.filterwarnings("ignore", message="Python 3.8 is no longer supported")

import os
import re
import shutil
import sys
import time

import openpyxl
import pandas as pd

from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import win32com.client
import win32com.client.gencache

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFileDialog, QRadioButton, QButtonGroup,
    QTextEdit, QProgressBar, QStackedWidget, QTreeWidget, QTreeWidgetItem,
    QMessageBox, QCheckBox
)
from PySide6.QtCore import Qt, QThread, Signal
from PySide6.QtGui import QFont

###################################################
# CONSTANTS
###################################################

__lenRevisionStr = 13       # len("Revision No: ")
__maxRetries = 3            # Word COM retry limit
__delay = 1                 # base seconds between retries

if getattr(sys, 'frozen', False):
    _SCRIPT_DIR = os.path.dirname(sys.executable)
else:
    _SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Default paths — overridden at runtime when user picks a folder
RECORDS_FOLDER  = os.path.join(_SCRIPT_DIR, 'KBM Docs')
OUTDATED_FOLDER = os.path.join(RECORDS_FOLDER, 'Outdated')
REDLINE_FOLDER  = os.path.join(RECORDS_FOLDER, 'Redline')
PREV_REV_FOLDER = os.path.join(RECORDS_FOLDER, 'Previous Revisions')

# DIA code -> document item number
KBM_MAP = {
    "DIA EF-732L.1": "101016",
    "DIA EF-732L.2": "101045",
    "DIA EF-732L.3": "101029",
    "DIA EF-733I.2": "103396",
    "DIA EF-736B.2": "101049",
    "DIA EF-736B.3": "101033",
    "DIA EF-736B.6": "102875",
    "DIA EF-737B.2": "101058",
    "DIA EF-737B.3": "101042",
    "DIA EF-737B.6": "102876",
    "DIA EF-738H": "103397",
    "DIA EF-753A.2": "103018",
    "DIA WF-7310": "101015",
    "DIA WF-732A": "100213",
    "DIA WF-732C": "100214",
    "DIA WF-732D": "100340",
    "DIA WF-732H": "101010",
    "DIA WF-733A.1": "100783",
    "DIA WF-733A.2": "101047",
    "DIA WF-733A.3": "101031",
    "DIA WF-733A.4": "101011",
    "DIA WF-733A.5": "103014",
    "DIA WF-733C": "100784",
    "DIA WF-733D": "100785",
    "DIA WF-733D.1": "100786",
    "DIA WF-734A.1": "101018",
    "DIA WF-734A.2": "101048",
    "DIA WF-734A.3": "101032",
    "DIA WF-735": "101012",
    "DIA WF-736B.1": "101019",
    "DIA WF-736B.4": "101050",
    "DIA WF-736B.5": "101034",
    "DIA WF-736E.1": "100864",
    "DIA WF-736E.2": "101051",
    "DIA WF-736E.3": "101035",
    "DIA WF-736G.1": "101024",
    "DIA WF-736G.2": "101052",
    "DIA WF-736G.3": "101036",
    "DIA WF-736H.1": "101025",
    "DIA WF-736H.2": "101053",
    "DIA WF-736H.3": "101037",
    "DIA WF-736J": "102786",
    "DIA WF-736L.1": "101026",
    "DIA WF-736L.2": "101054",
    "DIA WF-736L.3": "101038",
    "DIA WF-736N.2": "101055",
    "DIA WF-736N.3": "101039",
    "DIA WF-736N.4": "101056",
    "DIA WF-736N.5": "101040",
    "DIA WF-737A.1": "101027",
    "DIA WF-737A.2": "101057",
    "DIA WF-737A.3": "101041",
    "DIA WF-737B.1": "101028",
    "DIA WF-737B.4": "101059",
    "DIA WF-737B.5": "101043",
    "DIA WF-738A.1": "102429-001",
    "DIA WF-738A.2": "102427-001",
    "DIA WF-738A.3": "102430-000",
    "DIA WF-738G": "101013",
    "DIA WF-738H.1": "102545",
    "DIA WF-738H.2": "102546",
    "DIA WF-738H.3": "102547",
    "DIA WF-739A1": "101014",
    "DIA WF-753A": "102871",
    "DIA WF-753A.1": "102872",
    "DIA WF-753B.1": "102673",
    "DIA WF-754A": "102924",
    "DIA List of SOUP Items Project Whitney Web Application": "101046",
    "DIA List of SOUP Items Project Whitney Cloud Extensions": "101017",
    "DIA List of SOUP Items Project Whitney Mobile Application": "101030",
}

__lenVersionStr = 9         # len("Version: ") — offset to start of version value
__lenEndVersionStr = 14     # how far past "Version:" to read (captures "X.XX")

# General-purpose patterns (override the DIA-specific ones above for broader matching)
regexRevisionPattern = '[0-9]+\.[0-9]+'         # matches "1.2", "3.14"
regexDatePattern = '[0-9]+/[0-9]+/[0-9]+'       # matches "1/2/2024"
regexDocPathRevWhitespacePattern = r"[A-Za-z]+\s[0-9]+"
regexDocPathRev = r"[A-Za-z]+[0-9]+"

strRevCommentary = "External document updated."  # default revision history entry text
strVersion = ""

# Mutable state for the current processing run
revisionStr = ""
revisionInt = 0
docPathNew = ""

####################################################
# DOCUMENT PROCESSING FUNCTIONS
####################################################


# Read current rev from header, return the next one (A->B or 01->02)
def getRevisionInteger(section, strRev, newRev):
    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Revision" in cell.text:
                    strRev = cell.text[__lenRevisionStr:]

    if strRev.isalpha():
        newRev = chr(ord(strRev) + 1)
    else:
        newRev = int(strRev) + 1
        if newRev < 10:
            newRev = '0' + str(newRev)
        else:
            newRev = str(newRev)

    return newRev


# Bump the revision number in the DOCX header table
def headerRevUpdate(docObject, strRev, newRev):
    section = docObject.sections[0]
    newRev = getRevisionInteger(section, strRev, newRev)

    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Revision" in cell.text:
                    cell.text = "Revision No: " + newRev
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return newRev


# Append a new row to the revision history table at the end of the doc
def revHistoryUpdate(docObject, intRev, revHistoryTableCommentary, strVersion):
    listTables = docObject.tables
    if listTables:
        revHistoryTable = listTables[-1]
    else:
        revHistoryTable = None

    newRow = revHistoryTable.add_row()
    newRow.cells[0].text = str(intRev)
    newRow.cells[1].text = revHistoryTableCommentary + strVersion


# Save the doc with a new filename reflecting the updated rev
def docPathRevUpdate(docObject, oldDocPath, intRev, newDocPath):
    revEndFileNameWhitespaceMatch = re.search(regexDocPathRevWhitespacePattern, oldDocPath)

    if revEndFileNameWhitespaceMatch:
        newDocPath = re.sub(regexDocPathRevWhitespacePattern, ("Rev" + str(intRev)), oldDocPath)
    else:
        newDocPath = re.sub(regexDocPathRev, ("Rev" + str(intRev)), oldDocPath)

    docObject.save(newDocPath)
    return newDocPath


# Launch a COM app, clearing stale gen_py cache if needed
def handleDispatch(app_name):
    try:
        app = win32com.client.gencache.EnsureDispatch(app_name)
    except AttributeError:
        MODULE_LIST = [m.__name__ for m in sys.modules.values()]
        for module in MODULE_LIST:
            if re.match(r'win32com\.gen_py\..+', module):
                del sys.modules[module]
        shutil.rmtree(win32com.client.gencache.GetGeneratePath())
        app = win32com.client.gencache.EnsureDispatch(app_name)
    return app


# Wait until no WINWORD.EXE processes are running
def _wait_for_word_exit(timeout=30):
    import subprocess
    for _ in range(timeout):
        result = subprocess.run(
            ["tasklist", "/FI", "IMAGENAME eq WINWORD.EXE", "/NH"],
            capture_output=True, text=True,
            creationflags=subprocess.CREATE_NO_WINDOW
        )
        if "WINWORD.EXE" not in result.stdout:
            return
        time.sleep(1)


# Open both docs in Word and produce a _REDLINE.docx comparison
def createWordComparison(oldDocPath, newDocPath):
    numRetries = 0

    while numRetries < __maxRetries:
        _wait_for_word_exit()
        wordApp = None
        try:
            wordApp = handleDispatch("Word.Application")
            wordApp.Visible = False
            old_doc = wordApp.Documents.Open(oldDocPath, Visible=False)
            new_doc = wordApp.Documents.Open(newDocPath, Visible=False)
            wordApp.CompareDocuments(old_doc, new_doc)

        except Exception as e:
            print(f"[DEBUG] Attempt {numRetries + 1} failed: {e}")
            numRetries += 1
            if wordApp:
                try:
                    wordApp.Quit()
                except Exception:
                    pass
                del wordApp
                _wait_for_word_exit()
            time.sleep(__delay * numRetries)

        else:
            try:
                wordApp.ActiveDocument.SaveAs(FileName=(oldDocPath[:(len(oldDocPath) - 5)] + "_REDLINE.docx"))
                old_doc.Close(False)
                new_doc.Close(False)
                wordApp.Quit()
            except Exception:
                try:
                    wordApp.Quit()
                except Exception:
                    pass
            del wordApp
            _wait_for_word_exit()
            return

    raise Exception("Max retries exceeded. Comparison function not run.")


###################################################
# EXCEL UPDATE FUNCTIONS
###################################################


def normalize(s):
    return s.replace('_', ' ')


# Pull the "DIA XX-XXXX.X" code from a filename
def extract_dia_code(filename):
    normalized = normalize(filename)
    match = re.search(r'\bDIA\s+([A-Z0-9][A-Z0-9\-\.]*[A-Z0-9]|[A-Z]+)', normalized)
    if match:
        return f"DIA {match.group(1)}"
    return None


# Look up the document item number from KBM_MAP by DIA code
def get_id_from_json(data, dia_code, normalized_filename=None):
    entry = data.get(dia_code)
    if entry is not None:
        return entry if isinstance(entry, str) else entry.get('id')
    if normalized_filename:
        for key, value in data.items():
            if key in normalized_filename:
                return value if isinstance(value, str) else value.get('id')
    return None


# Extract (major, minor) version tuple from filename for sorting (e.g. "v1.3" -> (1,3))
def version_score(filename):
    m = re.search(r'\bv\.?(\d+)(?:\.(\d+))?', filename, re.IGNORECASE)
    if m:
        major = int(m.group(1))
        minor = int(m.group(2)) if m.group(2) else 0
        return (major, minor)
    return (0, 0)


# Walk folder for PDFs and DOCXs, skipping output subfolders
def collect_files(folder):
    pdf_files  = []
    docx_files = []
    skip = {'outdated', 'redline', 'previous revisions'}
    for root, dirs, files in os.walk(folder):
        dirs[:] = [d for d in dirs if d.lower() not in skip]
        for name in files:
            lower = name.lower()
            full  = os.path.abspath(os.path.join(root, name))
            if lower.endswith('.pdf'):
                pdf_files.append(full)
            elif lower.endswith('.docx'):
                docx_files.append(full)
    return pdf_files, docx_files


# Map 6-digit item number prefix -> list of DOCX paths for quick lookup
def build_docx_index(docx_files):
    index = {}
    pattern = re.compile(r'^(\d{6}(?:-\d{3,4})?)\b')
    for path in docx_files:
        name = os.path.basename(path)
        m = pattern.match(name)
        if m:
            prefix = m.group(1)
            index.setdefault(prefix, []).append(path)
    return index


def create_workbook():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['Index', 'PDF Path', 'DOCX Path'])
    ws.column_dimensions['B'].width = 124
    ws.column_dimensions['C'].width = 124
    return wb, ws


# Move older PDF versions to Outdated/, keeping only the highest-versioned one
def move_outdated_pdfs(all_pdfs_by_id, best_pdf):
    moved = 0
    os.makedirs(OUTDATED_FOLDER, exist_ok=True)
    for doc_id, entries in all_pdfs_by_id.items():
        if len(entries) <= 1:
            continue
        best_path = best_pdf[doc_id][1]
        for _, pdf_path in entries:
            if pdf_path == best_path:
                continue
            dest = os.path.join(OUTDATED_FOLDER, os.path.basename(pdf_path))
            shutil.move(pdf_path, dest)
            print(f'  [outdated] moved: {os.path.basename(pdf_path)}')
            moved += 1
    return moved


# Scan folder, match PDFs to DOCXs via KBM_MAP, write results to Excel
def update_excel(excel_path):
    data       = KBM_MAP
    pdf_files, docx_files = collect_files(RECORDS_FOLDER)

    docx_files = [p for p in docx_files if not re.search(r'red', os.path.basename(p), re.IGNORECASE)]
    docx_index = build_docx_index(docx_files)
    wb, ws     = create_workbook()
    start_row  = 2
    index      = 1

    no_dia  = 0
    no_json = 0

    all_pdfs_by_id = {}
    for pdf_path in pdf_files:
        name     = os.path.basename(pdf_path)
        dia_code = extract_dia_code(name)
        if dia_code is None:
            no_dia += 1
            continue
        doc_id = get_id_from_json(data, dia_code, normalize(name))
        if doc_id is None:
            no_json += 1
            print(f'  [no JSON match]  {name}  ->  {dia_code}')
            continue
        score = version_score(name)
        all_pdfs_by_id.setdefault(doc_id, []).append((score, pdf_path))

    best_pdf = {
        doc_id: max(entries, key=lambda x: x[0])
        for doc_id, entries in all_pdfs_by_id.items()
    }

    moved_count = move_outdated_pdfs(all_pdfs_by_id, best_pdf)

    no_docx        = 0
    matched        = 0
    missing_doc_ids = []

    for doc_id, (score, pdf_path) in sorted(best_pdf.items()):
        candidates = docx_index.get(doc_id, [])
        docx_path  = candidates[0] if candidates else None

        if docx_path is None:
            no_docx += 1
            missing_doc_ids.append((doc_id, os.path.basename(pdf_path)))

        ws.cell(row=start_row, column=1, value=index)
        ws.cell(row=start_row, column=2, value=pdf_path)
        ws.cell(row=start_row, column=3, value=docx_path if docx_path else '')

        index     += 1
        start_row += 1
        matched   += 1

    wb.save(excel_path)

    print(f'\nDone. Wrote {matched} rows to {excel_path}')
    print(f'  Outdated PDFs moved to outdated/: {moved_count}')
    print(f'  PDFs without DIA identifier:       {no_dia}')
    print(f'  DIA codes not in JSON:             {no_json}')
    print(f'  DIA codes with no DOCX match:      {no_docx}')

    if missing_doc_ids:
        print('\nDOCX files needed (add to folder):')
        for doc_id, pdf_name in missing_doc_ids:
            print(f'  {doc_id} -- {pdf_name}')

    return missing_doc_ids


####################################################
# DOCUMENT REVISION FUNCTIONS
####################################################


def extract_revision_from_filename(filename):
    match = re.search(r'Rev\s?([A-Z0-9]+)', filename, re.IGNORECASE)
    return match.group(1) if match else ""


# Sort key: letters first (A, B, C...) then numbers (01, 02...)
def revision_sort_key(rev_str):
    if rev_str.isalpha():
        return (0, rev_str)
    else:
        try:
            return (1, int(rev_str))
        except ValueError:
            return (2, rev_str)


# Find the two highest revisions of a doc for redline comparison
def find_top_two_revisions(docx_path):
    docx_basename = os.path.basename(docx_path)

    item_match = re.search(r"10\d{4}", docx_basename)
    if not item_match:
        return None, None

    item_no = item_match.group(0)

    # Always search from RECORDS_FOLDER root so nested or root-level docs are treated the same
    matching_files = []
    for root, _, files in os.walk(RECORDS_FOLDER):
        for fname in files:
            if fname.lower().endswith('.docx') and item_no in fname and 'red' not in fname.lower():
                full_path = os.path.join(root, fname)
                rev = extract_revision_from_filename(fname)
                matching_files.append((rev, full_path, fname))

    if len(matching_files) < 2:
        return None, None

    matching_files.sort(key=lambda x: revision_sort_key(x[0]))

    return matching_files[-2][1], matching_files[-1][1]


# Strip trailing blank paragraphs from the document
def remove_empty_pages(doc):
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)
        doc.paragraphs.pop()


# Extract the first date from the PDF's first page text
def find_date(firstPageStr):
    date_pattern = re.compile(
        r'\b(\d{1,2}/\d{1,2}/\d{4}|(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b',
        re.IGNORECASE
    )
    dates = date_pattern.findall(firstPageStr)
    return re.sub(r'\s+', ' ', dates[0][0]).strip()


# Read version and date from the PDF, open the matching DOCX
def read_word_document(pdf_path, docx_path):
    pdfReader = PdfReader(pdf_path)
    doc = Document(docx_path)
    page = pdfReader.pages[0]
    firstPageStr = page.extract_text()

    versionPos = firstPageStr.find("Version:") if firstPageStr.find("Version:") > -1 else firstPageStr.find("Version")
    versionStr = firstPageStr[(versionPos + __lenVersionStr):(versionPos + __lenEndVersionStr)].strip()
    dateStr     = find_date(firstPageStr)

    return doc, versionStr, dateStr


def process_documents(docDF, mode, doc_indices=None, progress_callback=None):
    global revisionStr, revisionInt, docPathNew

    if mode in (1, 3):
        os.makedirs(REDLINE_FOLDER, exist_ok=True)
    if mode in (1, 2):
        os.makedirs(PREV_REV_FOLDER, exist_ok=True)

    if doc_indices is None:
        rows_to_process = range(len(docDF))
    else:
        rows_to_process = doc_indices

    total = len(rows_to_process)
    pdf_paths = docDF["PDF Path"].values
    docx_paths = docDF["DOCX Path"].values

    print(f"Path: {RECORDS_FOLDER}\n")

    for i, row_idx in enumerate(rows_to_process):
        if progress_callback:
            progress_callback(i + 1, total)

        pdf = pdf_paths[row_idx]
        docPathInitial = docx_paths[row_idx]

        if not pdf or (isinstance(pdf, float) and pd.isna(pdf)) or not str(pdf).strip():
            print(f"[DEBUG] Skipped: Row {row_idx + 1} — Missing PDF path")
            continue
        if not docPathInitial or (isinstance(docPathInitial, float) and pd.isna(docPathInitial)) or not str(docPathInitial).strip():
            print(f"[DEBUG] Skipped: Row {row_idx + 1} — Missing DOCX path")
            continue

        print(f"\n{'='*60}")
        print(f"Processing: {os.path.basename(docPathInitial)}")

        if mode == 3:
            oldDocPath, newDocPath = find_top_two_revisions(docPathInitial)
            if oldDocPath is None or newDocPath is None:
                print(f"  [SKIPPED] Document has fewer than 2 versions. Cannot create redline.")
                continue

            print(f"  Comparing: {os.path.basename(oldDocPath)}")
            print(f"        vs:  {os.path.basename(newDocPath)}")
            createWordComparison(oldDocPath, newDocPath)

            redline_path = oldDocPath[:(len(oldDocPath) - 5)] + "_REDLINE.docx"
            if os.path.exists(redline_path):
                redline_filename = os.path.basename(redline_path)
                dest_redline_path = os.path.join(REDLINE_FOLDER, redline_filename)
                shutil.move(redline_path, dest_redline_path)
                print(f"Moved redline to Redline folder:{redline_filename}")

            continue

        doc, versionStr, dateStr = read_word_document(pdf, docPathInitial)

        revisionInt = headerRevUpdate(doc, revisionStr, revisionInt)

        textToUpdate = ""
        for para in doc.paragraphs:
            if "DIA" in para.text:
                textToUpdate = para.text

        textToUpdate = re.sub(regexRevisionPattern, versionStr, str(textToUpdate))
        textToUpdate = re.sub(regexDatePattern, dateStr, str(textToUpdate))

        for para in doc.paragraphs:
            if "DIA" in para.text:
                para.text = textToUpdate

        revHistoryUpdate(doc, revisionInt, strRevCommentary, strVersion)
        remove_empty_pages(doc)

        docPathNew = docPathRevUpdate(doc, docPathInitial, revisionInt, docPathNew)

        pdf_label = os.path.splitext(os.path.basename(pdf))[0]
        print(f"\nCreated: {os.path.basename(docPathNew)}")
        print(f"         [{pdf_label} – Date: {dateStr}]")

        if mode == 1:
            createWordComparison(docPathInitial, docPathNew)
            redline_path = docPathInitial[:(len(docPathInitial) - 5)] + "_REDLINE.docx"
            if os.path.exists(redline_path):
                dest_redline = os.path.join(REDLINE_FOLDER, os.path.basename(redline_path))
                shutil.move(redline_path, dest_redline)
                print(f"Moved redline to Redline folder:{os.path.basename(redline_path)}")
        else:
            print(f"[DEBUG] Skipped Redline: As requested for 'Create only new rev' mode")

        docx_filename = os.path.splitext(os.path.basename(docPathInitial))[0]

        item_match = re.search(r"10\d{4}", docx_filename)
        text_after = docx_filename[item_match.end():] if item_match else docx_filename
        text_after = re.sub(r"^[,\s]+", "", text_after)
        text_after = re.sub(r",?\s*Rev\s?[A-Z0-9]+", "", text_after).strip().rstrip(",").strip()

        dia_id = extract_dia_code(os.path.basename(pdf)) or ""

        part_number = item_match.group(0) if item_match else ""
        folder_prefix = f"{part_number}, " if part_number else ""
        # Always place organized folder at RECORDS_FOLDER root to avoid double-nesting on re-runs
        dest_folder = os.path.join(RECORDS_FOLDER, f"{folder_prefix}{text_after} ({dia_id})")

        if os.path.exists(docPathInitial) and not re.search(r'red', os.path.basename(docPathInitial), re.IGNORECASE):
            shutil.move(docPathInitial, os.path.join(PREV_REV_FOLDER, os.path.basename(docPathInitial)))
        if os.path.exists(docPathNew) and not re.search(r'red', os.path.basename(docPathNew), re.IGNORECASE):
            os.makedirs(dest_folder, exist_ok=True)
            shutil.move(docPathNew, os.path.join(dest_folder, os.path.basename(docPathNew)))
        if os.path.exists(pdf):
            os.makedirs(dest_folder, exist_ok=True)
            shutil.move(pdf, os.path.join(dest_folder, os.path.basename(pdf)))

        print(f"[DEBUG] Organized: {folder_prefix}{text_after} ({dia_id})")


####################################################
# GUI
####################################################


# Captures print() output and forwards it to a Qt signal for the log widget
class LogStream:
    def __init__(self, signal):
        self._signal = signal

    def write(self, text):
        if text and text.strip():
            self._signal.emit(text.rstrip('\n'))

    def flush(self):
        pass


# Background thread: scans folder and builds the Excel mapping
class ScanWorker(QThread):
    log_signal = Signal(str)
    finished_signal = Signal(object, object)
    error_signal = Signal(str)

    def __init__(self, excel_path):
        super().__init__()
        self._excel_path = excel_path

    def run(self):
        old_stdout = sys.stdout
        sys.stdout = LogStream(self.log_signal)
        try:
            missing_doc_ids = update_excel(self._excel_path)
            docDF = pd.read_excel(self._excel_path)
            self.finished_signal.emit(docDF, missing_doc_ids)
        except Exception as e:
            self.error_signal.emit(str(e))
        finally:
            sys.stdout = old_stdout


# Background thread: runs document revision/redline processing
class ProcessWorker(QThread):
    log_signal = Signal(str)
    progress_signal = Signal(int, int)
    finished_signal = Signal(bool, str)

    def __init__(self, docDF, mode, doc_indices):
        super().__init__()
        self._docDF = docDF
        self._mode = mode
        self._doc_indices = doc_indices

    def run(self):
        old_stdout = sys.stdout
        sys.stdout = LogStream(self.log_signal)
        try:
            process_documents(
                self._docDF,
                self._mode,
                self._doc_indices,
                progress_callback=lambda cur, total: self.progress_signal.emit(cur, total)
            )
            self.finished_signal.emit(True, "All operations completed successfully.")
        except Exception as e:
            self.finished_signal.emit(False, f"Error: {e}")
        finally:
            sys.stdout = old_stdout


# ─── Page 0: Folder Selection ────────────────────────────────────────────────


class FolderPage(QWidget):
    folder_changed = Signal(str)

    def __init__(self):
        super().__init__()
        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)

        layout = QVBoxLayout()
        layout.setSpacing(8)
        layout.setContentsMargins(16, 16, 16, 16)

        title = QLabel("Select KBM Docs Folder")
        title.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        layout.addWidget(title)

        desc = QLabel("Choose the folder containing your PDF and DOCX files.")
        desc.setWordWrap(True)
        layout.addWidget(desc)

        self.folder_label = QLabel("No folder selected")
        self.folder_label.setStyleSheet("color: #666; font-style: italic;")
        self.folder_label.setWordWrap(True)
        layout.addWidget(self.folder_label)

        browse_btn = QPushButton("Browse...")
        browse_btn.setFixedWidth(120)
        browse_btn.clicked.connect(self._browse)
        layout.addWidget(browse_btn)

        outer.addLayout(layout)
        outer.addStretch()

        self.selected_folder = ""

    def _browse(self):
        folder = QFileDialog.getExistingDirectory(self, "Select KBM Docs Folder")
        if folder:
            self.selected_folder = folder
            self.folder_label.setText(folder)
            self.folder_label.setStyleSheet("")
            self.folder_changed.emit(folder)

    def reset(self):
        self.selected_folder = ""
        self.folder_label.setText("No folder selected")
        self.folder_label.setStyleSheet("color: #666; font-style: italic;")


# ─── Page 1: Scan Results / Summary ──────────────────────────────────────────


class SummaryPage(QWidget):
    scan_finished = Signal()

    def __init__(self):
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        title = QLabel("File Summary")
        title.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        layout.addWidget(title)

        self.status_label = QLabel("Scanning folder...")
        layout.addWidget(self.status_label)

        self.summary_text = QTextEdit()
        self.summary_text.setReadOnly(True)
        self.summary_text.setFont(QFont("Consolas", 9))
        layout.addWidget(self.summary_text)

        self._last_scanned = None
        self._worker = None

    def on_enter(self, records_folder, excel_path):
        if records_folder == self._last_scanned:
            return

        if self._worker and self._worker.isRunning():
            return

        self.status_label.setText("Scanning folder...")
        self.summary_text.clear()

        global RECORDS_FOLDER, OUTDATED_FOLDER
        RECORDS_FOLDER = records_folder
        OUTDATED_FOLDER = os.path.join(RECORDS_FOLDER, 'Outdated')

        self._worker = ScanWorker(excel_path)
        self._worker.log_signal.connect(self._on_log)
        self._worker.finished_signal.connect(self._on_scan_done)
        self._worker.error_signal.connect(self._on_scan_error)
        self._worker.start()

    def _on_log(self, text):
        self.summary_text.append(text)

    def _on_scan_done(self, docDF, missing_doc_ids):
        self._last_scanned = RECORDS_FOLDER

        window = self.window()
        window.docDF = docDF
        window.missing_doc_ids = missing_doc_ids

        self.summary_text.append("\n" + "=" * 60)
        self.summary_text.append("FILE SUMMARY")
        self.summary_text.append("=" * 60)
        self._build_summary(docDF)

        matched_count = sum(
            1 for x in docDF["DOCX Path"].values
            if pd.notna(x) and str(x).strip()
        )
        missing_count = len(docDF) - matched_count
        self.status_label.setText(
            f"Scan complete: {matched_count} matched, {missing_count} missing"
        )

        self.scan_finished.emit()

    def _on_scan_error(self, error_msg):
        self.status_label.setText("Scan failed!")
        self.summary_text.append(f"\nERROR: {error_msg}")

    def _build_summary(self, docDF):
        pdf_paths = docDF["PDF Path"].values
        docx_paths = docDF["DOCX Path"].values

        matched = []
        missing = []

        for i in range(len(docDF)):
            pdf_name = os.path.basename(str(pdf_paths[i])) if pd.notna(pdf_paths[i]) else "N/A"
            docx_val = docx_paths[i]
            has_docx = pd.notna(docx_val) and str(docx_val).strip() != ""

            if has_docx:
                docx_name = os.path.basename(str(docx_val))
                matched.append((pdf_name, docx_name))
            else:
                missing.append(pdf_name)

        if matched:
            self.summary_text.append(f"\nMatched files ({len(matched)}):")
            for pdf_name, docx_name in matched:
                self.summary_text.append(f"\n  {pdf_name}\n  ->  {docx_name}")

        if missing:
            self.summary_text.append(f"\nMissing DOCX ({len(missing)}):")
            for pdf_name in missing:
                self.summary_text.append(f"  [!] {pdf_name}")

        self.summary_text.append(f"\nTotal: {len(matched)} matched, {len(missing)} missing")

    def reset(self):
        self.status_label.setText("Scanning folder...")
        self.summary_text.clear()
        self._last_scanned = None


# ─── Page 2: Missing Documents Warning ───────────────────────────────────────


class MissingDocsPage(QWidget):
    def __init__(self):
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setSpacing(12)

        title = QLabel("Missing Documents")
        title.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        layout.addWidget(title)

        desc = QLabel("The following DOCX files could not be found.\n"
                       "You can abort to fix your folder, or continue\n"
                       "(these documents will be skipped during processing).")
        desc.setWordWrap(True)
        layout.addWidget(desc)

        self.missing_list = QTextEdit()
        self.missing_list.setReadOnly(True)
        self.missing_list.setFont(QFont("Consolas", 9))
        layout.addWidget(self.missing_list)

    def on_enter(self, missing_doc_ids):
        lines = []
        for doc_id, pdf_name in missing_doc_ids:
            lines.append(f"  {doc_id}  --  {pdf_name}")
        self.missing_list.setText("\n".join(lines))


# ─── Page 3: Mode & Scope Selection ──────────────────────────────────────────


class ModePage(QWidget):
    def __init__(self):
        super().__init__()
        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)

        # Inner layout holds all controls packed tightly at the top
        layout = QVBoxLayout()
        layout.setSpacing(8)
        layout.setContentsMargins(16, 16, 16, 16)

        title = QLabel("Select Operation")
        title.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        layout.addWidget(title)

        # --- Mode selection ---
        mode_label = QLabel("Operation mode:")
        mode_label.setFont(QFont("Segoe UI", 10, QFont.Weight.Bold))
        layout.addWidget(mode_label)

        self.mode_group = QButtonGroup(self)
        self.mode_rev_and_redline = QRadioButton("Create new rev and redline")
        self.mode_rev_only = QRadioButton("Create only new rev")
        self.mode_redline_only = QRadioButton("Create only new redline")
        self.mode_rev_and_redline.setChecked(True)

        self.mode_group.addButton(self.mode_rev_and_redline, 1)
        self.mode_group.addButton(self.mode_rev_only, 2)
        self.mode_group.addButton(self.mode_redline_only, 3)

        layout.addWidget(self.mode_rev_and_redline)
        layout.addWidget(self.mode_rev_only)
        layout.addWidget(self.mode_redline_only)

        # --- Document checklist ---
        doc_list_label = QLabel("Documents to process:")
        doc_list_label.setFont(QFont("Segoe UI", 10, QFont.Weight.Bold))
        layout.addWidget(doc_list_label)

        sel_row = QHBoxLayout()
        self.select_all_cb = QCheckBox("Select All")
        self.select_new_cb = QCheckBox("Select New")
        sel_row.addWidget(self.select_all_cb)
        sel_row.addWidget(self.select_new_cb)
        sel_row.addStretch()
        layout.addLayout(sel_row)

        self.doc_list = QTreeWidget()
        self.doc_list.setHeaderHidden(True)
        self.doc_list.setRootIsDecorated(False)
        self.doc_list.setItemsExpandable(False)
        layout.addWidget(self.doc_list)

        outer.addLayout(layout)
        outer.addStretch()  # pushes everything up regardless of window height

        self.select_all_cb.stateChanged.connect(self._on_select_all_changed)
        self.select_new_cb.stateChanged.connect(self._on_select_new_changed)

    def _on_select_all_changed(self, state):
        check = Qt.CheckState.Checked if state == Qt.CheckState.Checked.value else Qt.CheckState.Unchecked
        for i in range(self.doc_list.topLevelItemCount()):
            top = self.doc_list.topLevelItem(i)
            if top.flags() & Qt.ItemFlag.ItemIsUserCheckable:
                top.setCheckState(0, check)
            for j in range(top.childCount()):
                child = top.child(j)
                if child.flags() & Qt.ItemFlag.ItemIsUserCheckable:
                    child.setCheckState(0, check)

    def _on_select_new_changed(self, state):
        check = Qt.CheckState.Checked if state == Qt.CheckState.Checked.value else Qt.CheckState.Unchecked
        for i in range(self.doc_list.topLevelItemCount()):
            top = self.doc_list.topLevelItem(i)
            if top.flags() & Qt.ItemFlag.ItemIsUserCheckable:
                top.setCheckState(0, check)

    def populate_docs(self, docDF):
        self.select_all_cb.setChecked(False)
        self.select_new_cb.setChecked(False)
        self.doc_list.clear()
        pdf_paths = docDF["PDF Path"].values

        root_docs = []
        grouped = {}

        for i, pdf_val in enumerate(pdf_paths):
            if not (pd.notna(pdf_val) and str(pdf_val).strip()):
                continue
            pdf_str = str(pdf_val)
            rel = os.path.relpath(os.path.dirname(pdf_str), RECORDS_FOLDER)
            entry = (i, pdf_str)
            if rel == '.':
                root_docs.append(entry)
            else:
                grouped.setdefault(rel, []).append(entry)

        def add_file_item(parent, row_idx, pdf_path):
            name = os.path.basename(pdf_path)
            dia_code = extract_dia_code(name)
            doc_id = get_id_from_json(KBM_MAP, dia_code, normalize(name)) if dia_code else None
            label = name + (f" ({doc_id})" if doc_id else "")
            item = QTreeWidgetItem(parent)
            item.setText(0, label)
            item.setFlags(Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsSelectable)
            item.setCheckState(0, Qt.CheckState.Unchecked)
            item.setData(0, Qt.ItemDataRole.UserRole, row_idx)

        for row_idx, pdf_path in root_docs:
            add_file_item(self.doc_list, row_idx, pdf_path)

        for folder_name, entries in sorted(grouped.items()):
            folder_item = QTreeWidgetItem(self.doc_list)
            folder_item.setText(0, folder_name)
            folder_item.setFlags(Qt.ItemFlag.ItemIsEnabled)
            folder_item.setFont(0, QFont("Segoe UI", 9, QFont.Weight.Bold))
            folder_item.setExpanded(True)
            for row_idx, pdf_path in entries:
                add_file_item(folder_item, row_idx, pdf_path)

    def get_mode(self):
        return self.mode_group.checkedId()

    def get_selected_indices(self):
        indices = []
        for i in range(self.doc_list.topLevelItemCount()):
            top = self.doc_list.topLevelItem(i)
            if (top.flags() & Qt.ItemFlag.ItemIsUserCheckable) and top.checkState(0) == Qt.CheckState.Checked:
                indices.append(top.data(0, Qt.ItemDataRole.UserRole))
            for j in range(top.childCount()):
                child = top.child(j)
                if (child.flags() & Qt.ItemFlag.ItemIsUserCheckable) and child.checkState(0) == Qt.CheckState.Checked:
                    indices.append(child.data(0, Qt.ItemDataRole.UserRole))
        return indices

    def reset(self):
        self.mode_rev_and_redline.setChecked(True)
        self.select_all_cb.setChecked(False)
        self.select_new_cb.setChecked(False)
        self.doc_list.clear()


# ─── Page 4: Processing / Results ────────────────────────────────────────────


class ProcessingPage(QWidget):
    def __init__(self):
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        self.title_label = QLabel("Processing Documents")
        self.title_label.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        layout.addWidget(self.title_label)

        self.progress_label = QLabel("")
        layout.addWidget(self.progress_label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

        self.log_output = QTextEdit()
        self.log_output.setReadOnly(True)
        self.log_output.setFont(QFont("Consolas", 9))
        layout.addWidget(self.log_output)

        self.debug_cb = QCheckBox("Show Debug Info")
        layout.addWidget(self.debug_cb)

        btn_layout = QHBoxLayout()
        self.save_btn = QPushButton("Save Log")
        self.start_over_btn = QPushButton("Start Over")
        self.close_btn = QPushButton("Close")

        self.save_btn.hide()
        self.start_over_btn.hide()
        self.close_btn.hide()

        btn_layout.addWidget(self.save_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(self.start_over_btn)
        btn_layout.addWidget(self.close_btn)
        layout.addLayout(btn_layout)

        self._log_entries = []  # list of (text, is_debug)

        self.save_btn.clicked.connect(self._save_log)
        self.close_btn.clicked.connect(QApplication.quit)
        self.debug_cb.stateChanged.connect(lambda _: self._render_log())

    def append_log(self, text):
        is_debug = text.strip().startswith("[DEBUG]")
        self._log_entries.append((text, is_debug))
        self._render_log()

    def _render_log(self):
        debug_on = self.debug_cb.isChecked()
        parts = []
        for text, is_debug in self._log_entries:
            if not is_debug or debug_on:
                parts.append(text if text.endswith("\n") else text + "\n")
        self.log_output.setPlainText("".join(parts))
        scrollbar = self.log_output.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    def on_progress(self, current, total):
        if total > 0:
            self.progress_bar.setMaximum(total)
            self.progress_bar.setValue(current)
            self.progress_label.setText(f"Document {current} of {total}")

    def on_finished(self, success, message):
        if success:
            self.title_label.setText("Complete")
        else:
            self.title_label.setText("Error")
        self.append_log(f"\n{'='*40}")
        self.append_log(message)
        self.progress_bar.hide()
        self.progress_label.hide()
        self._show_result_buttons()

    def show_aborted(self):
        self.title_label.setText("Aborted")
        self.progress_bar.hide()
        self.progress_label.hide()
        self._log_entries = []
        self.append_log("Operation aborted by user.")
        self.append_log("Missing DOCX files were detected and the user chose to abort.")
        self._show_result_buttons()

    def _show_result_buttons(self):
        self.save_btn.show()
        self.start_over_btn.show()
        self.close_btn.show()

    def _save_log(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Save Log", "whitney_log.txt", "Text Files (*.txt)"
        )
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(self.log_output.toPlainText())

    def reset(self):
        self.title_label.setText("Processing Documents")
        self.progress_label.setText("")
        self.progress_label.show()
        self.progress_bar.setValue(0)
        self.progress_bar.setMaximum(100)
        self.progress_bar.show()
        self._log_entries = []
        self.log_output.clear()
        self.save_btn.hide()
        self.start_over_btn.hide()
        self.close_btn.hide()


# ─── Main Window ─────────────────────────────────────────────────────────────


class WhitneyApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Whitney Document Tool")
        self.setMinimumSize(700, 500)
        self.resize(850, 620)

        self.records_folder = ""
        self.excel_path = ""
        self.docDF = None
        self.missing_doc_ids = []
        self._page_history = []
        self._process_worker = None

        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)

        self.stack = QStackedWidget()
        main_layout.addWidget(self.stack)

        # Navigation bar
        nav_layout = QHBoxLayout()
        self.back_btn = QPushButton("Back")
        self.abort_btn = QPushButton("Abort")
        self.next_btn = QPushButton("Next")

        self.abort_btn.setStyleSheet("color: #c0392b; font-weight: bold;")

        self.back_btn.clicked.connect(self._go_back)
        self.abort_btn.clicked.connect(self._on_abort)
        self.next_btn.clicked.connect(self._go_next)

        nav_layout.addWidget(self.back_btn)
        nav_layout.addStretch()
        nav_layout.addWidget(self.abort_btn)
        nav_layout.addWidget(self.next_btn)
        main_layout.addLayout(nav_layout)

        # Create pages
        self.folder_page = FolderPage()
        self.summary_page = SummaryPage()
        self.missing_page = MissingDocsPage()
        self.mode_page = ModePage()
        self.processing_page = ProcessingPage()

        self.stack.addWidget(self.folder_page)       # 0
        self.stack.addWidget(self.summary_page)      # 1
        self.stack.addWidget(self.missing_page)       # 2
        self.stack.addWidget(self.mode_page)          # 3
        self.stack.addWidget(self.processing_page)    # 4

        # Connections
        self.folder_page.folder_changed.connect(self._on_folder_changed)
        self.summary_page.scan_finished.connect(self._on_scan_finished)
        self.processing_page.start_over_btn.clicked.connect(self._start_over)

        self._update_nav()

    # ── Navigation ──

    def _on_folder_changed(self, folder):
        self.records_folder = folder
        self.excel_path = os.path.join(folder, "whitney_output.xlsx")
        self._update_nav()

    def _on_scan_finished(self):
        if self.missing_doc_ids:
            self._page_history.append(self.stack.currentIndex())
            self.missing_page.on_enter(self.missing_doc_ids)
            self._navigate_to(2)
        else:
            self._update_nav()

    def _go_next(self):
        current = self.stack.currentIndex()

        if current == 0:
            self._page_history.append(current)
            self._navigate_to(1)
            self.summary_page.on_enter(self.records_folder, self.excel_path)

        elif current == 1:
            self._page_history.append(current)
            self.mode_page.populate_docs(self.docDF)
            self._navigate_to(3)

        elif current == 2:
            self._page_history.append(current)
            self._navigate_to(1)

        elif current == 3:
            indices = self.mode_page.get_selected_indices()
            if indices is not None and len(indices) == 0:
                QMessageBox.warning(
                    self, "No Documents Selected",
                    "Please select at least one document to process."
                )
                return
            self._page_history.append(current)
            self._navigate_to(4)
            self._start_processing()

    def _go_back(self):
        if self._page_history:
            prev = self._page_history.pop()
            self._navigate_to(prev)

    def _on_abort(self):
        self._page_history.append(self.stack.currentIndex())
        self._navigate_to(4)
        self.processing_page.show_aborted()

    def _navigate_to(self, page_idx):
        self.stack.setCurrentIndex(page_idx)
        self._update_nav()

    def _update_nav(self):
        idx = self.stack.currentIndex()

        self.back_btn.setVisible(idx in (1, 2, 3))
        self.abort_btn.setVisible(idx == 2)

        if idx == 0:
            self.next_btn.setText("Next")
            self.next_btn.setEnabled(bool(self.records_folder))
            self.next_btn.show()
        elif idx == 1:
            self.next_btn.setText("Next")
            self.next_btn.setEnabled(self.docDF is not None)
            self.next_btn.show()
        elif idx == 2:
            self.next_btn.setText("Continue Anyway")
            self.next_btn.setEnabled(True)
            self.next_btn.show()
        elif idx == 3:
            self.next_btn.setText("Run")
            self.next_btn.setEnabled(True)
            self.next_btn.show()
        elif idx == 4:
            self.next_btn.hide()
            self.back_btn.hide()
            self.abort_btn.hide()

    # ── Processing ──

    def _start_processing(self):
        global RECORDS_FOLDER, REDLINE_FOLDER, PREV_REV_FOLDER, OUTDATED_FOLDER
        RECORDS_FOLDER = self.records_folder
        OUTDATED_FOLDER = os.path.join(RECORDS_FOLDER, 'Outdated')
        REDLINE_FOLDER = os.path.join(RECORDS_FOLDER, 'Redline')
        PREV_REV_FOLDER = os.path.join(RECORDS_FOLDER, 'Previous Revisions')

        self.processing_page.reset()

        mode = self.mode_page.get_mode()
        doc_indices = self.mode_page.get_selected_indices()

        self._process_worker = ProcessWorker(self.docDF, mode, doc_indices)
        self._process_worker.log_signal.connect(self.processing_page.append_log)
        self._process_worker.progress_signal.connect(self.processing_page.on_progress)
        self._process_worker.finished_signal.connect(self.processing_page.on_finished)
        self._process_worker.start()

    # ── Start Over / Close ──

    def _start_over(self):
        self._page_history.clear()
        self.records_folder = ""
        self.excel_path = ""
        self.docDF = None
        self.missing_doc_ids = []
        self.folder_page.reset()
        self.summary_page.reset()
        self.mode_page.reset()
        self.processing_page.reset()
        self._navigate_to(0)

    def closeEvent(self, event):
        if self._process_worker and self._process_worker.isRunning():
            reply = QMessageBox.question(
                self, "Processing in Progress",
                "Documents are still being processed. Are you sure you want to quit?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                event.ignore()
                return
        event.accept()


####################################################
# MAIN
####################################################


def main():
    app = QApplication(sys.argv)
    window = WhitneyApp()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
